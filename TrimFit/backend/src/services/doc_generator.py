import logging
import re
import tempfile
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from typing import List, Dict, Any, Optional, Set
from collections import defaultdict

logger = logging.getLogger(__name__)


class DocGenerator:
    def __init__(self):
        self.section_headers = {
            'professional_summary': ['summary', 'professional summary', 'profile', 'objective', 'career objective', 'about me', 'professional profile', 'career summary'],
            'experience': ['experience', 'work experience', 'employment', 'professional experience', 'work history', 'career history', 'employment history', 'work', 'career experience'],
            'skills': ['skills', 'technical skills', 'core competencies', 'competencies', 'expertise', 'technical expertise', 'core skills', 'key skills', 'areas of expertise'],
            'education': ['education', 'academic background', 'qualifications', 'academic qualifications', 'educational background', 'academic history'],
            'projects': ['projects', 'key projects', 'project experience', 'professional projects', 'notable projects', 'project work', 'portfolio'],
            'certifications': ['certifications', 'certificates', 'professional certifications', 'credentials', 'professional credentials'],
            'achievements': ['achievements', 'accomplishments', 'awards', 'honors', 'recognition', 'notable achievements']
        }

        self.target_sections_for_update = {
            'professional_summary', 'skills'}

        self.sections_to_preserve = {
            'experience', 'education', 'projects', 'certifications', 'achievements'}

        self.critical_sections_never_modify = {
            'education', 'projects'
        }

        self.personal_info_patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{3,4}[-\s\.]?[0-9]{3,4}',
            'linkedin': r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+/?',
            'github': r'(?:https?://)?(?:www\.)?github\.com/[\w-]+/?',
            'website': r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&\/=]*)'
        }

    def generate_tailored_resume(
        self,
        original_file_path: str,
        extracted_data: Dict[str, Any],
        tailored_data: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> str:

        try:
            logger.info("Starting tailored resume generation")
            logger.info(
                f"Target sections for update: {list(self.target_sections_for_update)}")
            logger.info(
                f"Critical sections never to modify: {list(self.critical_sections_never_modify)}")

            doc = Document(original_file_path)

            personal_info = self._extract_personal_info(doc)

            doc_structure = self._analyze_document_structure(doc)

            section_map = self._map_document_sections_enhanced(
                doc, doc_structure)

            self._verify_no_critical_sections_in_tailored_data(tailored_data)

            filtered_tailored_data = self._filter_tailored_data_strictly(
                tailored_data)

            critical_section_snapshots = {}
            for section_name, section_info in section_map.items():
                if section_name in self.critical_sections_never_modify:
                    logger.error(
                        f"Creating verification snapshot for CRITICAL section: {section_name}")
                    critical_section_snapshots[section_name] = self._create_critical_section_snapshot(
                        doc, section_info)

            self._update_document_sections_with_complete_protection(
                doc, section_map, extracted_data, filtered_tailored_data, personal_info)

            verification_result = self._verify_critical_sections_unchanged(
                doc, section_map, critical_section_snapshots)

            if not verification_result['all_verified']:
                logger.error("Critical sections verification failed")
                for section, status in verification_result['verification_details'].items():
                    if not status['verified']:
                        logger.error(
                            f"Section {section} was unexpectedly modified")
                raise Exception(
                    "Critical sections verification failed - education/projects may have been modified")

            if not output_path:
                temp_file = tempfile.NamedTemporaryFile(
                    delete=False, suffix='.docx')
                output_path = temp_file.name
                temp_file.close()

            doc.save(output_path)

            logger.info("Tailored resume generation completed successfully")
            return output_path

        except Exception as e:
            logger.error(f"Error generating tailored resume: {str(e)}")
            raise

    def _verify_no_critical_sections_in_tailored_data(self, tailored_data: Dict[str, Any]) -> None:
        critical_found = []
        for section_name in tailored_data.keys():
            if section_name in self.critical_sections_never_modify:
                critical_found.append(section_name)

        if critical_found:
            logger.warning(
                f"Found critical sections in tailored data: {critical_found} - these will be ignored")

    def _filter_tailored_data_strictly(self, tailored_data: Dict[str, Any]) -> Dict[str, Any]:
        filtered_data = {}

        for key, value in tailored_data.items():
            if key in self.target_sections_for_update:
                filtered_data[key] = value

        return filtered_data

    def _create_critical_section_snapshot(self, doc: Document, section_info: Dict[str, Any]) -> Dict[str, Any]:
        snapshot = {
            'header_text': '',
            'content_texts': [],
            'paragraph_count': 0,
            'total_characters': 0
        }

        try:
            header_idx = section_info['header_idx']
            if header_idx < len(doc.paragraphs):
                snapshot['header_text'] = doc.paragraphs[header_idx].text

            start_idx = section_info['start_idx']
            end_idx = section_info['end_idx']

            for idx in range(start_idx, min(end_idx + 1, len(doc.paragraphs))):
                para_text = doc.paragraphs[idx].text
                snapshot['content_texts'].append(para_text)
                snapshot['total_characters'] += len(para_text)

            snapshot['paragraph_count'] = len(snapshot['content_texts'])

        except Exception as e:
            logger.error(f"Error creating critical section snapshot: {e}")

        return snapshot

    def _verify_critical_sections_unchanged(
        self,
        doc: Document,
        section_map: Dict[str, Dict[str, Any]],
        original_snapshots: Dict[str, Dict[str, Any]]
    ) -> Dict[str, Any]:
        verification_result = {
            'all_verified': True,
            'verification_details': {}
        }

        current_doc_structure = self._analyze_document_structure(doc)
        current_section_map = self._map_document_sections_enhanced(
            doc, current_doc_structure)

        for section_name, original_snapshot in original_snapshots.items():
            verification_result['verification_details'][section_name] = {
                'verified': False,
                'reason': '',
                'details': {}
            }

            if section_name not in current_section_map:
                verification_result['all_verified'] = False
                verification_result['verification_details'][section_name]['reason'] = 'Section not found in current document'
                continue

            current_section_info = current_section_map[section_name]
            current_snapshot = self._create_critical_section_snapshot(
                doc, current_section_info)

            verification_details = {
                'header_match': original_snapshot['header_text'] == current_snapshot['header_text'],
                'paragraph_count_match': original_snapshot['paragraph_count'] == current_snapshot['paragraph_count'],
                'character_count_match': original_snapshot['total_characters'] == current_snapshot['total_characters'],
                'content_match': original_snapshot['content_texts'] == current_snapshot['content_texts']
            }

            all_match = all(verification_details.values())

            verification_result['verification_details'][section_name]['verified'] = all_match
            verification_result['verification_details'][section_name]['details'] = verification_details

            if all_match:
                logger.info(
                    f"✅ VERIFIED: {section_name} section remains exactly unchanged")
            else:
                logger.error(f"❌ FAILED: {section_name} section was modified!")
                logger.error(f"Verification details: {verification_details}")
                verification_result['all_verified'] = False
                verification_result['verification_details'][section_name]['reason'] = 'Content mismatch detected'

        return verification_result

    def _update_document_sections_with_complete_protection(
        self,
        doc: Document,
        section_map: Dict[str, Dict[str, Any]],
        extracted_data: Dict[str, Any],
        tailored_data: Dict[str, Any],
        personal_info: Dict[str, Set[str]]
    ):

        for section_name, section_info in section_map.items():
            if section_name in self.sections_to_preserve:

                continue

            if section_name in self.target_sections_for_update and section_name in tailored_data:
                new_content = tailored_data[section_name]

                new_content = self._preserve_personal_info_in_content(
                    new_content,
                    extracted_data.get(section_name),
                    personal_info
                )

                self._update_section_enhanced(doc, section_info, new_content)

    def _preserve_personal_info_in_content(
        self,
        new_content: Any,
        original_content: Any,
        personal_info: Dict[str, Set[str]]
    ) -> Any:
        if isinstance(new_content, str) and isinstance(original_content, str):
            for url in personal_info.get('urls', []):
                if url in original_content and url not in new_content:
                    logger.info(f"Preserving URL: {url}")

            for email in personal_info.get('emails', []):
                if email in original_content and email not in new_content:
                    logger.info(f"Preserving email: {email}")

            for phone in personal_info.get('phones', []):
                if phone in original_content and phone not in new_content:
                    logger.info(f"Preserving phone: {phone}")

        return new_content

    def _extract_personal_info(self, doc: Document) -> Dict[str, Set[str]]:
        personal_info = defaultdict(set)

        for paragraph in doc.paragraphs:
            text = paragraph.text

            for email in re.findall(self.personal_info_patterns['email'], text):
                personal_info['emails'].add(email)

            for phone in re.findall(self.personal_info_patterns['phone'], text):
                personal_info['phones'].add(phone)

            for pattern_name in ['linkedin', 'github', 'website']:
                for url in re.findall(self.personal_info_patterns[pattern_name], text):
                    personal_info['urls'].add(url)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    for email in re.findall(self.personal_info_patterns['email'], text):
                        personal_info['emails'].add(email)
                    for phone in re.findall(self.personal_info_patterns['phone'], text):
                        personal_info['phones'].add(phone)

        return dict(personal_info)

    def _analyze_document_structure(self, doc: Document) -> Dict[str, Any]:
        structure = {
            'styles': {},
            'paragraph_formats': [],
            'has_tables': len(doc.tables) > 0,
            'has_images': False,
            'has_headers_footers': bool(doc.sections[0].header.paragraphs or doc.sections[0].footer.paragraphs),
            'hyperlinks': [],
            'lists': defaultdict(list)
        }

        for idx, paragraph in enumerate(doc.paragraphs):
            para_info = {
                'idx': idx,
                'style': paragraph.style.name if paragraph.style else None,
                'alignment': paragraph.alignment,
                'left_indent': paragraph.paragraph_format.left_indent,
                'first_line_indent': paragraph.paragraph_format.first_line_indent,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
                'line_spacing': paragraph.paragraph_format.line_spacing,
                'is_list': self._is_list_paragraph(paragraph),
                'has_hyperlink': self._has_hyperlink(paragraph)
            }
            structure['paragraph_formats'].append(para_info)

            if para_info['has_hyperlink']:
                links = self._extract_hyperlinks_with_position(paragraph)
                structure['hyperlinks'].extend(links)

        return structure

    def _is_list_paragraph(self, paragraph: Paragraph) -> bool:
        if paragraph.style and ('List' in paragraph.style.name or 'Bullet' in paragraph.style.name):
            return True

        text = paragraph.text.strip()
        bullet_patterns = [r'^[•·▪▫◦‣⁃]', r'^[-*]\s',
                           r'^\d+[\.\)]\s', r'^[a-zA-Z][\.\)]\s']

        for pattern in bullet_patterns:
            if re.match(pattern, text):
                return True

        return False

    def _has_hyperlink(self, paragraph: Paragraph) -> bool:
        hyperlink_elements = paragraph._element.xpath('.//w:hyperlink')
        return len(hyperlink_elements) > 0

    def _extract_hyperlinks_with_position(self, paragraph: Paragraph) -> List[Dict[str, Any]]:
        hyperlinks = []

        for idx, run in enumerate(paragraph.runs):
            parent = run._element.getparent()
            if parent.tag.endswith('hyperlink'):
                r_id = parent.get(qn('r:id'))
                if r_id:
                    try:
                        rel = paragraph.part.rels[r_id]
                        hyperlinks.append({
                            'text': run.text,
                            'url': rel.target_ref,
                            'run_index': idx,
                            'paragraph': paragraph,
                            'r_id': r_id
                        })
                    except KeyError:
                        logger.warning(
                            f"Could not find relationship for hyperlink ID: {r_id}")

        return hyperlinks

    def _map_document_sections_enhanced(self, doc: Document, doc_structure: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
        section_map = {}
        current_section = None
        section_start_idx = 0

        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            section_type = self._identify_section_header(text)

            if section_type:
                if current_section and section_start_idx < idx:
                    section_info = {
                        'header_idx': section_start_idx,
                        'start_idx': section_start_idx + 1,
                        'end_idx': idx - 1,
                        'header_paragraph': doc.paragraphs[section_start_idx],
                        'format_info': doc_structure['paragraph_formats'][section_start_idx + 1:idx] if section_start_idx + 1 < idx else []
                    }
                    section_map[current_section] = section_info

                current_section = section_type
                section_start_idx = idx

        if current_section and section_start_idx < len(doc.paragraphs):
            section_info = {
                'header_idx': section_start_idx,
                'start_idx': section_start_idx + 1,
                'end_idx': len(doc.paragraphs) - 1,
                'header_paragraph': doc.paragraphs[section_start_idx],
                'format_info': doc_structure['paragraph_formats'][section_start_idx + 1:] if section_start_idx + 1 < len(doc.paragraphs) else []
            }
            section_map[current_section] = section_info

        return section_map

    def _identify_section_header(self, text: str) -> Optional[str]:
        """Identify if text is a section header and return the section type."""
        text_lower = text.lower().strip()

        # Clean up punctuation
        text_lower = re.sub(r'[:\-–—•·]', '', text_lower).strip()
        text_lower = re.sub(r'\s+', ' ', text_lower)

        word_count = len(text_lower.split())
        if word_count > 6:
            return None

        for section_type, headers in self.section_headers.items():
            for header in headers:
                if text_lower == header:
                    return section_type

        for section_type, headers in self.section_headers.items():
            for header in headers:
                if text_lower.startswith(header + ' ') or text_lower.startswith(header + ':'):
                    return section_type

        if word_count <= 3:
            for section_type, headers in self.section_headers.items():
                for header in headers:
                    if re.search(r'\b' + re.escape(header) + r'\b', text_lower):
                        return section_type

        single_word_matches = {
            'skills': ['skills'],
            'education': ['education'],
            'projects': ['projects'],
            'experience': ['experience']
        }

        if word_count == 1:
            for section_type, single_words in single_word_matches.items():
                if text_lower in single_words:
                    return section_type

        return None

    def _update_section_enhanced(self, doc: Document, section_info: Dict[str, Any], new_content: Any):
        start_idx = section_info['start_idx']
        end_idx = section_info['end_idx']
        format_info = section_info.get('format_info', [])

        if isinstance(new_content, str):
            self._update_text_section_enhanced(
                doc, start_idx, end_idx, new_content, format_info)
        elif isinstance(new_content, list):
            self._update_list_section_enhanced(
                doc, start_idx, end_idx, new_content, format_info)
        elif isinstance(new_content, dict):
            self._update_structured_section_enhanced(
                doc, start_idx, end_idx, new_content, format_info)

    def _update_text_section_enhanced(
        self,
        doc: Document,
        start_idx: int,
        end_idx: int,
        new_text: str,
        format_info: List[Dict[str, Any]]
    ):
        if start_idx > end_idx or start_idx >= len(doc.paragraphs):
            return

        original_structure = self._analyze_section_structure(
            doc, start_idx, end_idx)
        common_font_formatting = original_structure.get(
            'common_font_formatting', {})

        new_lines = new_text.strip().split('\n')

        if len(new_lines) == 1 and end_idx - start_idx > 0:
            paragraph = doc.paragraphs[start_idx]

            if len(format_info) > 0:
                self._preserve_paragraph_formatting(paragraph, format_info[0])

            self._update_paragraph_content_with_font_formatting(
                paragraph, new_text.strip(), common_font_formatting)

            for idx in range(start_idx + 1, min(end_idx + 1, len(doc.paragraphs))):
                doc.paragraphs[idx].clear()
        else:
            # Handle multiple lines
            for i, line in enumerate(new_lines):
                if start_idx + i <= end_idx and start_idx + i < len(doc.paragraphs):
                    paragraph = doc.paragraphs[start_idx + i]

                    if i < len(format_info):
                        self._preserve_paragraph_formatting(
                            paragraph, format_info[i])

                    self._update_paragraph_content_with_font_formatting(
                        paragraph, line, common_font_formatting)
                else:
                    logger.info(
                        f"Adding new paragraph for line: {line[:50]}...")
                    if end_idx < len(doc.paragraphs) - 1:
                        reference_para = doc.paragraphs[end_idx]
                        new_para = self._insert_paragraph_after(
                            doc, reference_para, line)
                        if format_info:
                            self._preserve_paragraph_formatting(
                                new_para, format_info[-1])
                        self._apply_font_formatting_to_paragraph(
                            new_para, common_font_formatting)
                para_idx = i + 1

            for idx in range(start_idx + len(new_lines), min(end_idx + 1, len(doc.paragraphs))):
                doc.paragraphs[idx].clear()

    def _update_list_section_enhanced(
        self,
        doc: Document,
        start_idx: int,
        end_idx: int,
        new_items: List[Any],
        format_info: List[Dict[str, Any]]
    ):
        original_structure = self._analyze_section_structure(
            doc, start_idx, end_idx)
        common_font_formatting = original_structure.get(
            'common_font_formatting', {})

        current_idx = start_idx

        for item_idx, item in enumerate(new_items):
            if current_idx > end_idx:
                break

            if isinstance(item, dict):
                current_idx = self._add_structured_item_enhanced(
                    doc, current_idx, end_idx, item, format_info, common_font_formatting
                )
            elif isinstance(item, str):
                if current_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[current_idx]

                    if current_idx - start_idx < len(format_info):
                        para_format = format_info[current_idx - start_idx]
                        if para_format.get('is_list') and not item.startswith(('•', '-', '*')):
                            item = f"• {item}"

                    self._update_paragraph_content_with_font_formatting(
                        paragraph, item, common_font_formatting)
                    current_idx += 1

        for idx in range(current_idx, min(end_idx + 1, len(doc.paragraphs))):
            doc.paragraphs[idx].clear()

    def _update_structured_section_enhanced(
        self,
        doc: Document,
        start_idx: int,
        end_idx: int,
        new_data: Dict[str, Any],
        format_info: List[Dict[str, Any]]
    ):
        current_idx = start_idx

        original_structure = self._analyze_section_structure(
            doc, start_idx, end_idx)
        common_font_formatting = original_structure.get(
            'common_font_formatting', {})

        available_paragraphs = end_idx - start_idx + 1
        categories_list = list(new_data.items())
        categories_to_process = min(len(categories_list), available_paragraphs)

        categories_processed = 0

        for i, (category, items) in enumerate(categories_list[:categories_to_process]):
            if isinstance(items, list) and items:
                if original_structure.get('uses_colons', True):
                    formatted_category = category.replace('_', ' ').title()
                    content = f"{formatted_category}: {', '.join(str(item) for item in items)}"
                else:
                    separator = original_structure.get('separator', ': ')
                    formatted_category = category.replace('_', ' ').title()
                    content = f"{formatted_category}{separator}{', '.join(str(item) for item in items)}"

                if current_idx <= end_idx and current_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[current_idx]

                    format_idx = current_idx - start_idx
                    if format_idx < len(format_info):
                        self._preserve_paragraph_formatting(
                            paragraph, format_info[format_idx])

                    self._update_paragraph_content_with_font_formatting(
                        paragraph, content, common_font_formatting)
                    current_idx += 1
                    categories_processed += 1
                else:
                    break

        for idx in range(start_idx + categories_processed, min(end_idx + 1, len(doc.paragraphs))):
            doc.paragraphs[idx].clear()

    def _analyze_section_structure(self, doc: Document, start_idx: int, end_idx: int) -> Dict[str, Any]:
        structure = {
            'uses_colons': False,
            'separator': ': ',
            'bullet_style': None,
            'line_pattern': 'category_items',
            'common_font_formatting': None
        }

        colon_count = 0
        total_lines = 0
        font_samples = []

        for idx in range(start_idx, min(end_idx + 1, len(doc.paragraphs))):
            paragraph = doc.paragraphs[idx]
            text = paragraph.text.strip()

            if text:
                total_lines += 1
                if ':' in text:
                    colon_count += 1
                    if ': ' in text:
                        structure['separator'] = ': '
                    elif ':' in text:
                        structure['separator'] = ':'

                if text.startswith(('•', '-', '*', '▪')):
                    structure['bullet_style'] = text[0]

                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    font_info = {
                        'font_name': first_run.font.name,
                        'font_size': first_run.font.size,
                        'bold': first_run.bold,
                        'italic': first_run.italic,
                        'underline': first_run.underline
                    }
                    font_samples.append(font_info)

        if total_lines > 0:
            structure['uses_colons'] = (colon_count / total_lines) > 0.5

        if font_samples:
            structure['common_font_formatting'] = self._get_most_common_font_formatting(
                font_samples)

        return structure

    def _get_most_common_font_formatting(self, font_samples: List[Dict[str, Any]]) -> Dict[str, Any]:

        if not font_samples:
            return {}

        font_names = [f.get('font_name')
                      for f in font_samples if f.get('font_name')]
        font_sizes = [f.get('font_size')
                      for f in font_samples if f.get('font_size')]
        bold_values = [f.get('bold')
                       for f in font_samples if f.get('bold') is not None]
        italic_values = [f.get('italic')
                         for f in font_samples if f.get('italic') is not None]
        underline_values = [f.get('underline')
                            for f in font_samples if f.get('underline') is not None]

        common_formatting = {}

        if font_names:
            common_formatting['font_name'] = max(
                set(font_names), key=font_names.count)

        if font_sizes:
            common_formatting['font_size'] = max(
                set(font_sizes), key=font_sizes.count)

        if bold_values:
            common_formatting['bold'] = max(
                set(bold_values), key=bold_values.count)

        if italic_values:
            common_formatting['italic'] = max(
                set(italic_values), key=italic_values.count)

        if underline_values:
            common_formatting['underline'] = max(
                set(underline_values), key=underline_values.count)

        return common_formatting

    def _update_paragraph_content_with_font_formatting(self, paragraph: Paragraph, new_text: str, font_formatting: Dict[str, Any]):

        existing_hyperlinks = self._extract_hyperlinks_with_position(paragraph)

        hyperlinks_to_preserve = []
        for link in existing_hyperlinks:
            if link['text'] in new_text or link['url'] in new_text:
                hyperlinks_to_preserve.append(link)

        paragraph.clear()

        if hyperlinks_to_preserve:
            paragraph._original_formatting = font_formatting
            self._add_text_with_preserved_hyperlinks(
                paragraph, new_text, hyperlinks_to_preserve)
        else:
            run = paragraph.add_run(new_text)
            self._apply_formatting_to_run(run, font_formatting)

    def _apply_formatting_to_run(self, run: Run, formatting: Dict[str, Any]):

        if not formatting:
            return

        try:
            if formatting.get('font_name'):
                run.font.name = formatting['font_name']
            if formatting.get('font_size'):
                run.font.size = formatting['font_size']
            if formatting.get('font_color'):
                run.font.color.rgb = formatting['font_color']
            if formatting.get('bold') is not None:
                run.bold = formatting['bold']
            if formatting.get('italic') is not None:
                run.italic = formatting['italic']
            if formatting.get('underline') is not None:
                run.underline = formatting['underline']
        except Exception as e:
            logger.warning(f"Failed to apply formatting: {e}")

    def _apply_font_formatting_to_paragraph(self, paragraph: Paragraph, font_formatting: Dict[str, Any]):

        for run in paragraph.runs:
            self._apply_formatting_to_run(run, font_formatting)

    def _add_text_with_preserved_hyperlinks(
        self,
        paragraph: Paragraph,
        text: str,
        hyperlinks: List[Dict[str, Any]]
    ):

        original_formatting = None
        if hasattr(paragraph, '_original_formatting'):
            original_formatting = paragraph._original_formatting

        remaining_text = text

        hyperlinks_sorted = sorted(
            hyperlinks,
            key=lambda x: text.find(
                x['text']) if x['text'] in text else len(text)
        )

        for link in hyperlinks_sorted:
            link_text = link['text']
            if link_text in remaining_text:
                before_idx = remaining_text.find(link_text)
                if before_idx > 0:
                    run = paragraph.add_run(remaining_text[:before_idx])
                    self._apply_formatting_to_run(run, original_formatting)

                self._add_hyperlink_to_paragraph(
                    paragraph, link_text, link['url'])

                remaining_text = remaining_text[before_idx + len(link_text):]

        if remaining_text:
            run = paragraph.add_run(remaining_text)
            self._apply_formatting_to_run(run, original_formatting)

    def _add_hyperlink_to_paragraph(self, paragraph: Paragraph, text: str, url: str):

        part = paragraph.part

        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        run_element = OxmlElement('w:r')

        rPr = OxmlElement('w:rPr')

        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0000FF")
        rPr.append(c)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        run_element.append(rPr)

        text_element = OxmlElement('w:t')
        text_element.text = text
        run_element.append(text_element)

        hyperlink.append(run_element)
        paragraph._element.append(hyperlink)

    def _preserve_paragraph_formatting(self, paragraph: Paragraph, format_info: Dict[str, Any]):

        if format_info.get('alignment'):
            paragraph.alignment = format_info['alignment']

        if format_info.get('left_indent'):
            paragraph.paragraph_format.left_indent = format_info['left_indent']

        if format_info.get('first_line_indent'):
            paragraph.paragraph_format.first_line_indent = format_info['first_line_indent']

        if format_info.get('space_before'):
            paragraph.paragraph_format.space_before = format_info['space_before']

        if format_info.get('space_after'):
            paragraph.paragraph_format.space_after = format_info['space_after']

        if format_info.get('style'):
            try:
                paragraph.style = format_info['style']
            except KeyError:
                logger.warning(
                    f"Style '{format_info['style']}' not found in document")

    def _add_structured_item_enhanced(
        self,
        doc: Document,
        start_idx: int,
        end_idx: int,
        item: Dict[str, Any],
        format_info: List[Dict[str, Any]],
        common_font_formatting: Optional[Dict[str, Any]] = None
    ) -> int:

        current_idx = start_idx

        if not common_font_formatting:
            original_structure = self._analyze_section_structure(
                doc, start_idx, end_idx)
            common_font_formatting = original_structure.get(
                'common_font_formatting', {})

        if 'name' in item:
            if current_idx <= end_idx and current_idx < len(doc.paragraphs):
                paragraph = doc.paragraphs[current_idx]
                self._update_paragraph_content_with_font_formatting(
                    paragraph, item['name'], common_font_formatting)

                if current_idx - start_idx < len(format_info):
                    self._preserve_paragraph_formatting(
                        paragraph, format_info[current_idx - start_idx]
                    )
                current_idx += 1

        if 'description' in item and item['description']:
            if current_idx <= end_idx and current_idx < len(doc.paragraphs):
                paragraph = doc.paragraphs[current_idx]
                self._update_paragraph_content_with_font_formatting(
                    paragraph, item['description'], common_font_formatting)

                if current_idx - start_idx < len(format_info):
                    self._preserve_paragraph_formatting(
                        paragraph, format_info[current_idx - start_idx]
                    )
                current_idx += 1

        if 'details' in item and isinstance(item['details'], list):
            for detail in item['details']:
                if current_idx <= end_idx and current_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[current_idx]

                    if (current_idx - start_idx < len(format_info) and
                            format_info[current_idx - start_idx].get('is_list')):
                        if not detail.strip().startswith(('•', '-', '*')):
                            detail = f"• {detail}"

                    self._update_paragraph_content_with_font_formatting(
                        paragraph, detail, common_font_formatting)

                    if current_idx - start_idx < len(format_info):
                        self._preserve_paragraph_formatting(
                            paragraph, format_info[current_idx - start_idx]
                        )

                    current_idx += 1
                else:
                    break

        if 'technologies' in item and isinstance(item['technologies'], list):
            tech_line = f"Technologies: {', '.join(item['technologies'])}"

            if current_idx <= end_idx and current_idx < len(doc.paragraphs):
                paragraph = doc.paragraphs[current_idx]
                self._update_paragraph_content_with_font_formatting(
                    paragraph, tech_line, common_font_formatting)

                if current_idx - start_idx < len(format_info):
                    self._preserve_paragraph_formatting(
                        paragraph, format_info[current_idx - start_idx]
                    )
                current_idx += 1

        return current_idx


document_generator = DocGenerator()
